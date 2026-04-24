import os
import sys
import time
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import lyricsgenius
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

API_TOKEN = "jiUdC-afCCPNZ5VxXwvJ9F0lgoQQhM07D8ddFHly9CtHe3Rp3Sz5ba6Y9tLKN6uJ"
INPUT_FILE = "songs.txt"
OUTPUT_DIR = "output"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "歌词合集.docx")

def parse_songs_file(filepath):
    songs = []
    if not os.path.exists(filepath):
        print(f"错误：找不到文件 {filepath}")
        return songs
    
    with open(filepath, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            
            if ' - ' in line:
                parts = line.split(' - ', 1)
                title = parts[0].strip()
                artist = parts[1].strip() if len(parts) > 1 else ""
            else:
                title = line
                artist = ""
            
            songs.append({"title": title, "artist": artist})
    
    return songs

def search_lyrics(genius, title, artist=""):
    try:
        if artist:
            song = genius.search_song(title, artist)
        else:
            search_results = genius.search_songs(title)
            if search_results and len(search_results) > 0:
                song = search_results[0]
            else:
                return None, None, None
        
        if song is None:
            return None, None, None
        
        song_title = song.title
        song_artist = song.artist
        
        lyrics = getattr(song, 'lyrics', None) or getattr(song, '_body', None)
        
        return song_title, song_artist, lyrics
    
    except Exception as e:
        print(f"  搜索出错: {e}")
        return None, None, None

def create_word_document(songs_data, output_path):
    doc = Document()
    
    title_paragraph = doc.add_heading('歌词合集', 0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph("目录")
    doc.add_paragraph()
    
    for i, song in enumerate(songs_data, 1):
        doc.add_paragraph(f"{i}. {song['title']} - {song['artist']}")
    
    for song in songs_data:
        add_page_break(doc)
        
        heading = doc.add_heading(f"{song['title']} - {song['artist']}", level=1)
        
        if song['lyrics']:
            para = doc.add_paragraph(song['lyrics'])
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
        else:
            doc.add_paragraph("（未找到歌词）")
    
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"\n文档已保存至: {output_path}")

def add_page_break(doc):
    doc.add_page_break()

def main():
    print("=" * 50)
    print("歌词批量下载工具 - Word 文档生成器")
    print("=" * 50)
    
    songs = parse_songs_file(INPUT_FILE)
    if not songs:
        print("没有找到任何歌曲，请检查 songs.txt 文件")
        return
    
    print(f"已加载 {len(songs)} 首歌曲\n")
    
    genius = lyricsgenius.Genius(API_TOKEN)
    genius.verbose = False
    genius.retries = 3
    
    songs_data = []
    
    for i, song in enumerate(songs, 1):
        print(f"[{i}/{len(songs)}] 正在获取: {song['title']}", end="")
        if song['artist']:
            print(f" - {song['artist']}", end="")
        print()
        
        title, artist, lyrics = search_lyrics(genius, song['title'], song['artist'])
        
        if lyrics:
            songs_data.append({
                "title": title or song['title'],
                "artist": artist or song['artist'] or "未知艺术家",
                "lyrics": lyrics
            })
            print(f"  ✓ 成功获取歌词")
        else:
            songs_data.append({
                "title": song['title'],
                "artist": song['artist'] or "未知艺术家",
                "lyrics": None
            })
            print(f"  ✗ 未找到歌词")
        
        time.sleep(0.5)
    
    if songs_data:
        create_word_document(songs_data, OUTPUT_FILE)
        print("\n完成！")
    else:
        print("\n未能获取任何歌词")

if __name__ == "__main__":
    main()