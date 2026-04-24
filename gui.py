import sys
import os
import time
import io
import re
import requests
import logging
from datetime import datetime
from bs4 import BeautifulSoup
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                             QHBoxLayout, QLabel, QLineEdit, QPushButton, 
                             QProgressBar, QTextEdit, QFileDialog, QMessageBox,
                             QGroupBox, QDialog, QCheckBox, QScrollArea, QFrame,
                             QListWidget, QListWidgetItem, QDialogButtonBox)
from PyQt6.QtCore import QThread, pyqtSignal, Qt
from PyQt6.QtGui import QFont, QColor, QTextCharFormat, QSyntaxHighlighter

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# 日志配置
LOG_DIR = "output"
LOG_FILE = os.path.join(LOG_DIR, "app.log")

os.makedirs(LOG_DIR, exist_ok=True)

file_handler = logging.FileHandler(LOG_FILE, encoding='utf-8', mode='a')
file_handler.setLevel(logging.DEBUG)
file_handler.setFormatter(logging.Formatter('%(asctime)s [%(levelname)s] %(message)s'))

logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[file_handler]
)
logger = logging.getLogger(__name__)

import lyricsgenius
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

API_TOKEN = "jiUdC-afCCPNZ5VxXwvJ9F0lgoQQhM07D8ddFHly9CtHe3Rp3Sz5ba6Y9tLKN6uJ"

def fetch_lyrics_from_azlyrics(title, artist):
    logger.info(f"尝试获取歌词: {title} - {artist}")
    try:
        # First try lyrics.ovh API (more reliable)
        search_artist = artist.split(',')[0].split('&')[0].split('ft')[0].split('feat')[0].strip()
        search_title = title.split('(')[0].strip()
        
        url = f"https://api.lyrics.ovh/v1/{search_artist}/{search_title}"
        logger.debug(f"lyrics.ovh URL: {url}")
        
        headers = {
            'User-Agent': 'Mozilla/5.0',
        }
        
        time.sleep(0.5)
        
        response = requests.get(url, headers=headers, timeout=15)
        logger.debug(f"lyrics.ovh 响应状态码: {response.status_code}")
        
        if response.status_code == 200:
            data = response.json()
            lyrics = data.get('lyrics', '').strip()
            if lyrics and len(lyrics) > 50:
                return lyrics, f"{title} - {artist} (lyrics.ovh)"
        
        # Fallback: try AZLyrics (may be blocked)
        artist_clean = re.sub(r'[^a-zA-Z0-9]', '', artist.lower()).strip()
        title_clean = re.sub(r'[^a-zA-Z0-9]', '', title.lower()).strip()
        
        if not artist_clean:
            artist_clean = re.sub(r'[^a-zA-Z0-9]', '', title.lower()).strip()
            title_clean = ""
        
        url = f"https://www.azlyrics.com/lyrics/{artist_clean}/{title_clean}.html"
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
        }
        
        time.sleep(1)
        
        response = requests.get(url, headers=headers, timeout=15)
        
        if response.status_code != 200:
            return None, f"HTTP {response.status_code}"
        
        if 'unusual activity' in response.text.lower() or 'checking your browser' in response.text.lower():
            return None, "lyrics.ovh 失败，AZLyrics 被封禁"
        
        logger.warning(f"lyrics.ovh 请求被封禁: {title} - {artist}")
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        content_div = soup.find('div', class_=lambda x: x and 'text-center' in x and 'col-xs-12' in x)
        
        if not content_div:
            return None, "lyrics.ovh 失败，找不到 AZLyrics 内容"
        
        logger.warning(f"AZLyrics 内容为空: {title} - {artist}")
        
        text = content_div.get_text()
        
        if len(text) < 500:
            return None, f"lyrics.ovh 失败，AZLyrics 内容太短"
        
        lines = text.split('\n')
        cleaned_lines = []
        start_found = False
        
        for line in lines:
            line = line.strip()
            line_lower = line.lower()
            
            if 'usage of azlyrics' in line_lower:
                start_found = True
                continue
            if not start_found:
                continue
            if 'mxm' in line_lower:
                break
            if line and len(line) > 1:
                cleaned_lines.append(line)
        
        lyrics = '\n'.join(cleaned_lines).strip()
        
        if len(lyrics) > 50:
            logger.info(f"AZLyrics 成功获取歌词: {title} - {artist}, 长度={len(lyrics)}")
            return lyrics, f"{title} - {artist} (AZLyrics)"
        
        logger.warning(f"lyrics.ovh 和 AZLyrics 都失败: {title} - {artist}")
        return None, "lyrics.ovh 和 AZLyrics 都失败"
    
    except Exception as e:
        logger.error(f"fetch_lyrics_from_azlyrics 异常: {title} - {artist}, 错误={str(e)}")
        return None, str(e)

def fetch_lyrics_from_netease(title, artist):
    logger.info(f"尝试网易云: {title} - {artist}")
    try:
        search_title = title.split('(')[0].strip()
        search_artist = artist.split(',')[0].split('&')[0].strip() if artist else ""
        
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
            "Referer": "http://music.163.com",
            "Host": "music.163.com"
        }
        
        # 方案1：同时使用歌曲名和艺术家搜索
        search_query = f"{search_title} {search_artist}".strip() if search_artist else search_title
        search_url = "http://music.163.com/api/search/get/"
        search_data = {"s": search_query, "type": 1, "limit": 10, "offset": 0}
        
        logger.debug(f"网易云搜索: {search_query}")
        
        time.sleep(0.5)
        
        response = requests.post(search_url, data=search_data, headers=headers, timeout=15)
        logger.debug(f"网易云搜索响应: {response.status_code}")
        
        if response.status_code != 200:
            logger.error(f"网易云搜索失败 HTTP {response.status_code}: {title} - {artist}")
            return None, f"搜索失败: HTTP {response.status_code}"
        
        result = response.json()
        songs = result.get('result', {}).get('songs', [])
        
        if not songs:
            logger.warning(f"网易云未找到歌曲: {title}")
            return None, "未找到歌曲"
        
        # 方案2：遍历结果，查找匹配艺术家的版本
        matched_song = None
        if search_artist:
            for s in songs:
                s_artist = s['artists'][0]['name'] if s.get('artists') else ""
                # 检查艺术家名是否包含在搜索的艺术家名中，或者反过来
                if search_artist.lower() in s_artist.lower() or s_artist.lower() in search_artist.lower():
                    matched_song = s
                    break
        
        # 如果没有匹配的，使用第一个结果
        if not matched_song:
            matched_song = songs[0]
        
        song_id = matched_song['id']
        song_name = matched_song['name']
        artist_name = matched_song['artists'][0]['name']
        
        logger.debug(f"网易云找到歌曲: {song_name} - {artist_name}, id={song_id}")
        
        lyrics_url = f"http://music.163.com/api/song/lyric?id={song_id}&lv=-1&tv=-1"
        
        response = requests.get(lyrics_url, headers=headers, timeout=15)
        
        if response.status_code != 200:
            logger.error(f"网易云获取歌词失败 HTTP {response.status_code}")
            return None, f"获取歌词失败: HTTP {response.status_code}"
        
        lyrics_result = response.json()
        
        lrc = lyrics_result.get('lrc', {}).get('lyric', '')
        
        if not lrc:
            logger.warning(f"网易云无歌词: {title} - {artist_name}")
            return None, "无歌词"
        
        # 移除时间标签 [00:00.45]
        lyrics = re.sub(r'\[\d{2}:\d{2}\.\d{2,3}\]', '', lrc)
        lyrics = re.sub(r'\n\n+', '\n', lyrics).strip()
        
        logger.info(f"网易云成功获取歌词: {title} - {artist_name}, 长度={len(lyrics)}")
        return lyrics, f"{song_name} - {artist_name} (网易云)"
    
    except Exception as e:
        logger.error(f"网易云异常: {title} - {artist}, 错误={str(e)}")
        return None, str(e)


class LyricsDownloaderThread(QThread):
    progress_updated = pyqtSignal(int, int, str)
    log_message = pyqtSignal(str)
    finished_with_result = pyqtSignal(list)
    stopped = pyqtSignal()

    def __init__(self, input_file, output_file):
        super().__init__()
        self.input_file = input_file
        self.output_file = output_file
        self.is_running = True
        self.songs_data = []

    def parse_songs_file(self, filepath):
        songs = []
        if not os.path.exists(filepath):
            return songs
        
        with open(filepath, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                
                if ' - ' in line:
                    parts = line.split(' - ', 1)
                    title = parts[0].strip()
                    artist_raw = parts[1].strip() if len(parts) > 1 else ""
                    
                    if '/' in artist_raw:
                        artist = artist_raw.split('/')[0].strip()
                    elif ',' in artist_raw:
                        artist = artist_raw.split(',')[0].strip()
                    else:
                        artist = artist_raw
                else:
                    title = line
                    artist = ""
                
                songs.append({"title": title, "artist": artist})
        
        return songs

    def normalize_text(self, text):
        if not text:
            return ""
        text = text.lower().strip()
        text = text.replace('&', 'and')
        text = text.replace('ft.', 'feat.')
        text = text.replace('ft ', 'feat ')
        text = text.replace('feat.', 'feat.')
        text = text.replace('featuring', 'feat.')
        text = text.replace("'", "'").replace("'", "'").replace("'", "'")
        text = text.replace("'", "'").replace("'", "'").replace("'", "'")
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace('...', '')
        import re
        text = re.sub(r'\s*\(.*?\)', '', text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def fuzzy_match(self, input_text, matched_text):
        if not input_text or not matched_text:
            return False
        input_norm = self.normalize_text(input_text)
        matched_norm = self.normalize_text(matched_text)
        if not input_norm or not matched_norm:
            return False
        return input_norm in matched_norm or matched_norm in input_norm

    def check_match(self, input_title, input_artist, matched_title, matched_artist):
        title_match = self.fuzzy_match(input_title, matched_title)
        
        if not input_artist:
            if title_match:
                return "ok", "匹配正确"
            else:
                return "error", "歌名不匹配"
        
        artist_match = self.fuzzy_match(input_artist, matched_artist)
        
        if title_match and artist_match:
            return "ok", "匹配正确"
        elif title_match and not artist_match:
            return "warning", f"艺术家可能不同: {matched_artist}"
        elif not title_match and artist_match:
            return "warning", f"歌名可能不同: {matched_title}"
        else:
            return "error", f"不匹配: {matched_title} - {matched_artist}"

    def search_lyrics(self, genius, title, artist=""):
        try:
            if artist:
                song = genius.search_song(title, artist)
            else:
                search_results = genius.search_songs(title)
                if search_results and len(search_results) > 0:
                    song = search_results[0]
                else:
                    return None, None, None
            
            if song is None:
                return None, None, None
            
            song_title = song.title
            song_artist = song.artist
            lyrics = getattr(song, 'lyrics', None) or getattr(song, '_body', None)
            
            return song_title, song_artist, lyrics
        
        except Exception as e:
            self.log_message.emit(f"搜索出错: {e}")
            return None, None, None

    def run(self):
        logger.info("=" * 40)
        logger.info("程序启动 - 开始搜索歌词")
        
        self.log_message.emit("=" * 40)
        self.log_message.emit("开始搜索歌词...")
        
        try:
            songs = self.parse_songs_file(self.input_file)
            if not songs:
                logger.error(f"找不到歌曲列表文件: {self.input_file}")
                self.log_message.emit("错误：找不到歌曲列表文件")
                self.stopped.emit()
                return
            
            logger.info(f"加载歌曲列表成功: {len(songs)} 首")
            self.log_message.emit(f"已加载 {len(songs)} 首歌曲")
            
            logger.info("初始化 Genius API...")
            genius = lyricsgenius.Genius(API_TOKEN)
            genius.verbose = False
            genius.retries = 3
            logger.info("Genius API 初始化完成")
            
            songs_data = []
            
            for i, song in enumerate(songs, 1):
                if not self.is_running:
                    logger.info("用户停止下载")
                    self.log_message.emit("已停止")
                    self.stopped.emit()
                    return
                
                logger.info(f"[{i}/{len(songs)}] 开始搜索: {song['title']} - {song['artist']}")
                self.progress_updated.emit(i, len(songs), f"{song['title']} - {song['artist']}")
                
                title, artist, lyrics = self.search_lyrics(genius, song['title'], song['artist'])
                
                if lyrics:
                    status, message = self.check_match(song['title'], song['artist'], title, artist)
                    logger.info(f"[{status.upper()}] 搜索成功: {title} - {artist} ({message})")
                    songs_data.append({
                        "input_title": song['title'],
                        "input_artist": song['artist'],
                        "matched_title": title,
                        "matched_artist": artist,
                        "lyrics": lyrics,
                        "status": status,
                        "check_message": message
                    })
                    self.log_message.emit(f"[{status.upper()}] [{i}/{len(songs)}] {title} - {artist} ({message})")
                else:
                    logger.warning(f"未找到歌词: {song['title']} - {song['artist']}")
                    songs_data.append({
                        "input_title": song['title'],
                        "input_artist": song['artist'],
                        "matched_title": song['title'],
                        "matched_artist": song['artist'],
                        "lyrics": None,
                        "status": "error",
                        "check_message": "未找到歌词"
                    })
                    self.log_message.emit(f"[ERROR] [{i}/{len(songs)}] 未找到歌词: {song['title']}")
                
                time.sleep(0.5)
            
            logger.info(f"歌词获取完成，共 {len(songs_data)} 首")
            self.log_message.emit(f"\n歌词获取完成，共 {len(songs_data)} 首")
            self.finished_with_result.emit(songs_data)
            
        except Exception as e:
            logger.error(f"下载线程异常: {str(e)}", exc_info=True)
            self.log_message.emit(f"下载出错: {str(e)}")
            self.stopped.emit()

    def stop(self):
        self.is_running = False


class PreviewDialog(QDialog):
    def __init__(self, songs_data, output_file, parent=None):
        super().__init__(parent)
        self.songs_data = songs_data
        self.output_file = output_file
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("检查匹配结果")
        self.setMinimumSize(700, 500)
        self.current_theme = "light"
        
        layout = QVBoxLayout()
        
        title_label = QLabel("请检查以下歌曲的匹配结果：")
        title_label.setStyleSheet("font-size: 14pt; font-weight: bold;")
        layout.addWidget(title_label)
        
        info_label = QLabel("⚠️ = 需确认  |  ❌ = 错误  |  ✅ = 匹配正确（已隐藏）")
        info_label.setStyleSheet("color: #666;")
        layout.addWidget(info_label)
        
        layout.addSpacing(10)
        
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll_widget = QWidget()
        scroll_widget.setObjectName("scrollWidget")
        scroll_layout = QVBoxLayout()
        
        for i, song in enumerate(self.songs_data):
            if song["status"] in ("pending", "warning", "error", "searching"):
                item_frame = self.create_item_frame(i, song)
                scroll_layout.addWidget(item_frame)
        
        if scroll_layout.count() == 0:
            no_issues_label = QLabel("所有歌曲已处理完成！")
            no_issues_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            no_issues_label.setStyleSheet("font-size: 12pt; color: green; padding: 20px;")
            scroll_layout.addWidget(no_issues_label)
        
        scroll_layout.addStretch()
        scroll_widget.setLayout(scroll_layout)
        scroll.setWidget(scroll_widget)
        
        self.scroll_layout = scroll_layout
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setFixedHeight(25)
        layout.addWidget(self.progress_bar)
        
        layout.addWidget(scroll)
        
        button_layout = QHBoxLayout()
        
        self.theme_btn = QPushButton("切换主题")
        self.theme_btn.clicked.connect(self.toggle_theme)
        button_layout.addWidget(self.theme_btn)
        
        self.select_all_btn = QPushButton("全选")
        self.select_all_btn.clicked.connect(lambda: self.set_all_checked(True))
        button_layout.addWidget(self.select_all_btn)
        
        self.deselect_all_btn = QPushButton("取消全选")
        self.deselect_all_btn.clicked.connect(lambda: self.set_all_checked(False))
        button_layout.addWidget(self.deselect_all_btn)
        
        self.retry_btn = QPushButton("一键重试")
        self.retry_btn.setObjectName("retryBtn")
        self.retry_btn.setStyleSheet("""
            QPushButton#retryBtn { 
                background-color: #e91e63; 
                color: white; 
                font-size: 12pt; 
                padding: 10px 20px;
                border-radius: 6px;
            }
        """)
        self.retry_btn.clicked.connect(self.retry_all_checked)
        button_layout.addWidget(self.retry_btn)
        
        button_layout.addStretch()
        
        self.confirm_btn = QPushButton("生成文档")
        self.confirm_btn.setObjectName("confirmBtn")
        self.confirm_btn.setStyleSheet("""
            QPushButton#confirmBtn { 
                background-color: #4CAF50; 
                color: white; 
                font-size: 12pt; 
                padding: 10px 20px;
                border-radius: 6px;
            }
        """)
        self.confirm_btn.clicked.connect(self.confirm_and_generate)
        button_layout.addWidget(self.confirm_btn)
        
        layout.addLayout(button_layout)
        
        self.setLayout(layout)
        self.apply_theme()
    
    def apply_theme(self):
        if self.current_theme == "dark":
            self.setStyleSheet("""
                QDialog { background-color: #1e1e1e; color: #e0e0e0; }
                QLabel { color: #e0e0e0; }
                QFrame { 
                    background-color: #2d2d2d; 
                    border: 1px solid #444; 
                    border-radius: 8px;
                    margin: 4px;
                }
                QCheckBox { 
                    color: #e0e0e0; 
                    font-size: 12pt;
                    padding: 8px;
                }
                QCheckBox::indicator {
                    width: 18px;
                    height: 18px;
                    border-radius: 4px;
                    background-color: #3d3d3d;
                    border: 2px solid #555;
                }
                QCheckBox::indicator:checked {
                    background-color: #4CAF50;
                    border-color: #4CAF50;
                }
                QCheckBox:hover {
                    background-color: #3d3d3d;
                    border-radius: 4px;
                }
                QPushButton { 
                    background-color: #0d47a1; 
                    color: white; 
                    border: none;
                    border-radius: 6px;
                    padding: 10px 20px;
                    font-size: 11pt;
                }
                QPushButton:hover { background-color: #1565c0; }
                QPushButton#confirmBtn {
                    background-color: #2e7d32;
                }
                QPushButton#confirmBtn:hover { background-color: #388e3c; }
                QScrollArea { 
                    background-color: #1e1e1e; 
                    border: none;
                }
                QWidget#scrollWidget { background-color: #1e1e1e; }
            """)
        else:
            self.setStyleSheet("""
                QDialog { background-color: #fafafa; color: #333; }
                QLabel { color: #333; }
                QFrame { 
                    background-color: #ffffff; 
                    border: 1px solid #e0e0e0; 
                    border-radius: 8px;
                    margin: 4px;
                }
                QCheckBox { 
                    color: #333; 
                    font-size: 12pt;
                    padding: 8px;
                }
                QCheckBox::indicator {
                    width: 18px;
                    height: 18px;
                    border-radius: 4px;
                    background-color: #f5f5f5;
                    border: 2px solid #ddd;
                }
                QCheckBox::indicator:checked {
                    background-color: #4CAF50;
                    border-color: #4CAF50;
                }
                QCheckBox:hover {
                    background-color: #f0f0f0;
                    border-radius: 4px;
                }
                QPushButton { 
                    background-color: #2196F3; 
                    color: white; 
                    border: none;
                    border-radius: 6px;
                    padding: 10px 20px;
                    font-size: 11pt;
                }
                QPushButton:hover { background-color: #1976D2; }
                QPushButton#confirmBtn {
                    background-color: #4CAF50;
                }
                QPushButton#confirmBtn:hover { background-color: #388E3C; }
                QScrollArea { 
                    background-color: #fafafa; 
                    border: none;
                }
                QWidget#scrollWidget { background-color: #fafafa; }
            """)

    def toggle_theme(self):
        if self.current_theme == "light":
            self.current_theme = "dark"
        else:
            self.current_theme = "light"
        self.apply_theme()
    
    def create_item_frame(self, index, song):
        frame = QFrame()
        frame.setFrameStyle(QFrame.Shape.Box)
        frame.setLineWidth(1)
        
        main_layout = QVBoxLayout()
        main_layout.setContentsMargins(10, 8, 10, 8)
        
        status_icon = "✅"
        status_text = "匹配正确"
        if song["status"] == "warning":
            status_icon = "⚠️"
            status_text = song["check_message"]
        elif song["status"] == "error":
            status_icon = "❌"
            status_text = song["check_message"]
        elif song["status"] == "pending":
            status_icon = "⏳"
            status_text = song.get("check_message", "待确认")
        elif song["status"] == "searching":
            status_icon = "🔍"
            status_text = song.get("check_message", "正在搜索...")
        elif song["status"] == "skipped":
            status_icon = "⏭️"
            status_text = song.get("check_message", "跳过")
        
        checkbox = QCheckBox(f"[{status_icon}] {index + 1}. {song['input_title']} - {song['input_artist']}")
        checkbox.setChecked(song["status"] in ("pending", "warning", "error", "searching"))
        checkbox.setStyleSheet("font-size: 11pt;")
        checkbox.song_index = index
        
        main_layout.addWidget(checkbox)
        
        info_label = QLabel(f"({status_text})")
        info_label.setStyleSheet("font-size: 9pt; color: #666; margin-left: 20px;")
        main_layout.addWidget(info_label)
        
        if song["status"] in ("pending", "warning", "error", "searching"):
            button_layout = QHBoxLayout()
            button_layout.setContentsMargins(20, 5, 0, 0)
            
            retry_btn = QPushButton("重新获取")
            retry_btn.setFixedWidth(100)
            retry_btn.setStyleSheet("""
                QPushButton {
                    background-color: #ff9800;
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 6px;
                    font-size: 10pt;
                }
                QPushButton:hover { background-color: #f57c00; }
            """)
            retry_btn.song_index = index
            retry_btn.clicked.connect(lambda: self.retry_from_any_source(index))
            button_layout.addWidget(retry_btn)
            
            netease_btn = QPushButton("网易云")
            netease_btn.setFixedWidth(80)
            netease_btn.setStyleSheet("""
                QPushButton {
                    background-color: #e91e63;
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 6px;
                    font-size: 10pt;
                }
                QPushButton:hover { background-color: #c2185b; }
            """)
            netease_btn.song_index = index
            netease_btn.clicked.connect(lambda: self.retry_from_netease(index))
            button_layout.addWidget(netease_btn)
            
            skip_btn = QPushButton("跳过")
            skip_btn.setFixedWidth(80)
            skip_btn.setStyleSheet("""
                QPushButton {
                    background-color: #9e9e9e;
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 6px;
                    font-size: 10pt;
                }
                QPushButton:hover { background-color: #757575; }
            """)
            skip_btn.song_index = index
            skip_btn.clicked.connect(lambda: self.skip_song(index))
            button_layout.addWidget(skip_btn)
            
            confirm_btn = QPushButton("确认")
            confirm_btn.setFixedWidth(80)
            confirm_btn.setStyleSheet("""
                QPushButton {
                    background-color: #2196F3;
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 6px;
                    font-size: 10pt;
                }
                QPushButton:hover { background-color: #1976D2; }
            """)
            confirm_btn.song_index = index
            confirm_btn.clicked.connect(lambda: self.confirm_status_only(index))
            button_layout.addWidget(confirm_btn)
            
            button_layout.addStretch()
            
            main_layout.addLayout(button_layout)
        
        frame.setLayout(main_layout)
        frame.checkbox = checkbox
        return frame
    
    def retry_from_azlyrics(self, index):
        song = self.songs_data[index]
        title = song["input_title"]
        artist = song["input_artist"]
        
        btn = self.sender()
        if btn:
            btn.setEnabled(False)
            btn.setText("获取中...")
        
        lyrics, msg = fetch_lyrics_from_azlyrics(title, artist)
        
        if lyrics:
            self.songs_data[index]["lyrics"] = lyrics
            self.songs_data[index]["matched_title"] = title
            self.songs_data[index]["matched_artist"] = artist
            self.songs_data[index]["status"] = "ok"
            self.songs_data[index]["check_message"] = "获取成功"
            
            QMessageBox.information(self, "成功", f"已获取歌词:\n{title} - {artist}")
            self.accept()
            
            from PyQt6.QtWidgets import QApplication
            QApplication.instance().processEvents()
            
            new_dialog = PreviewDialog(self.songs_data, self.output_file, self.parent())
            new_dialog.exec()
        else:
            QMessageBox.warning(self, "失败", f"获取歌词失败:\n{msg}")
            if btn:
                btn.setEnabled(True)
                btn.setText("重新获取")
    
    def retry_from_netease(self, index):
        song = self.songs_data[index]
        title = song["input_title"]
        artist = song["input_artist"]
        
        btn = self.sender()
        if btn:
            btn.setEnabled(False)
            btn.setText("获取中...")
        
        self.songs_data[index]["status"] = "searching"
        self.songs_data[index]["check_message"] = "正在搜索..."
        self.refresh_list()
        
        lyrics, msg = fetch_lyrics_from_netease(title, artist)
        
        if lyrics:
            matched_title = title
            matched_artist = artist
            if " - " in msg:
                parts = msg.split(" - ", 1)
                matched_title = parts[0].strip()
                matched_artist = parts[1].strip()
            
            self.songs_data[index]["lyrics"] = lyrics
            self.songs_data[index]["matched_title"] = matched_title
            self.songs_data[index]["matched_artist"] = matched_artist
            self.songs_data[index]["status"] = "pending"
            self.songs_data[index]["check_message"] = msg
            
            QMessageBox.information(self, "成功", f"已从网易云获取歌词:\n{msg}\n请确认")
            self.refresh_list()
        else:
            self.songs_data[index]["status"] = "skipped"
            self.songs_data[index]["check_message"] = f"获取失败: {msg}"
            QMessageBox.warning(self, "失败", f"网易云获取失败:\n{msg}")
            self.refresh_list()
    
    def retry_from_any_source(self, index):
        song = self.songs_data[index]
        title = song["input_title"]
        artist = song["input_artist"]
        
        btn = self.sender()
        if btn:
            btn.setEnabled(False)
            btn.setText("获取中...")
        
        self.songs_data[index]["status"] = "searching"
        self.songs_data[index]["check_message"] = "正在搜索..."
        self.refresh_list()
        
        lyrics, msg = fetch_lyrics_from_azlyrics(title, artist)
        
        if not lyrics:
            lyrics, msg = fetch_lyrics_from_netease(title, artist)
        
        if lyrics:
            matched_title = title
            matched_artist = artist
            if " - " in msg:
                parts = msg.split(" - ", 1)
                matched_title = parts[0].strip()
                matched_artist = parts[1].strip()
            
            self.songs_data[index]["lyrics"] = lyrics
            self.songs_data[index]["matched_title"] = matched_title
            self.songs_data[index]["matched_artist"] = matched_artist
            self.songs_data[index]["status"] = "pending"
            self.songs_data[index]["check_message"] = msg
            
            QMessageBox.information(self, "成功", f"已获取歌词:\n{msg}\n请确认")
            self.refresh_list()
        else:
            self.songs_data[index]["status"] = "skipped"
            self.songs_data[index]["check_message"] = f"获取失败: {msg}"
            QMessageBox.warning(self, "失败", f"获取歌词失败:\n{msg}")
            self.refresh_list()
    
    def skip_song(self, index):
        logger.info(f"用户跳过歌曲: {self.songs_data[index]['input_title']}")
        self.songs_data[index]["status"] = "skipped"
        self.songs_data[index]["check_message"] = "用户跳过"
        self.refresh_list()
    
    def confirm_status_only(self, index):
        logger.info(f"用户确认歌曲: {self.songs_data[index]['input_title']}")
        self.songs_data[index]["status"] = "ok"
        self.songs_data[index]["check_message"] = "用户确认包含"
        self.refresh_list()
    
    def retry_all_checked(self):
        checked_indices = []
        for i in range(self.scroll_layout.count()):
            widget = self.scroll_layout.itemAt(i).widget()
            if widget and hasattr(widget, 'checkbox'):
                if widget.checkbox.isChecked():
                    song_index = widget.checkbox.song_index
                    song = self.songs_data[song_index]
                    if song["status"] in ("pending", "warning", "error", "searching"):
                        checked_indices.append(song_index)
        
        if not checked_indices:
            QMessageBox.warning(self, "警告", "请先勾选需要重新搜索的歌曲")
            return
        
        total = len(checked_indices)
        self.progress_bar.setVisible(True)
        self.progress_bar.setMaximum(total)
        self.progress_bar.setValue(0)
        self.progress_bar.setFormat(f"正在搜索 0/{total}...")
        
        self.retry_btn.setEnabled(False)
        
        from PyQt6.QtWidgets import QApplication
        
        for idx, song_index in enumerate(checked_indices):
            song = self.songs_data[song_index]
            title = song["input_title"]
            artist = song["input_artist"]
            
            self.songs_data[song_index]["status"] = "searching"
            self.songs_data[song_index]["check_message"] = "正在搜索..."
            self.progress_bar.setFormat(f"正在搜索 {idx+1}/{total}...")
            QApplication.instance().processEvents()
            
            lyrics, msg = fetch_lyrics_from_netease(title, artist)
            
            if lyrics:
                matched_title = title
                matched_artist = artist
                if " - " in msg:
                    parts = msg.split(" - ", 1)
                    matched_title = parts[0].strip()
                    matched_artist = parts[1].strip()
                
                self.songs_data[song_index]["lyrics"] = lyrics
                self.songs_data[song_index]["matched_title"] = matched_title
                self.songs_data[song_index]["matched_artist"] = matched_artist
                self.songs_data[song_index]["status"] = "pending"
                self.songs_data[song_index]["check_message"] = msg
            else:
                self.songs_data[song_index]["status"] = "skipped"
                self.songs_data[song_index]["check_message"] = f"获取失败: {msg}"
            
            self.progress_bar.setValue(idx + 1)
            QApplication.instance().processEvents()
        
        self.progress_bar.setVisible(False)
        self.retry_btn.setEnabled(True)
        
        self.refresh_list()
        
        success_count = sum(1 for i in checked_indices if self.songs_data[i]["status"] == "pending")
        QMessageBox.information(self, "完成", f"搜索完成：{success_count}/{total} 首成功")
    
    def find_checkbox_for_song(self, index):
        for i in range(self.scroll_layout.count()):
            widget = self.scroll_layout.itemAt(i).widget()
            if widget and hasattr(widget, 'checkbox') and hasattr(widget.checkbox, 'song_index'):
                if widget.checkbox.song_index == index:
                    return widget.checkbox
        return None
    
    def refresh_list(self):
        logger.debug("刷新预览列表")
        
        # 先移除stretch（最后一个item）
        if self.scroll_layout.count() > 0:
            last_item = self.scroll_layout.itemAt(self.scroll_layout.count() - 1)
            if last_item and last_item.spacerItem():
                self.scroll_layout.takeAt(self.scroll_layout.count() - 1)
        
        # 移除所有现有的widget
        while self.scroll_layout.count() > 0:
            item = self.scroll_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        
        # 按原始顺序添加pending的widget
        has_pending = False
        for i, song in enumerate(self.songs_data):
            if song["status"] in ("pending", "warning", "error", "searching"):
                has_pending = True
                item_frame = self.create_item_frame(i, song)
                self.scroll_layout.addWidget(item_frame)
        
        if not has_pending:
            no_issues_label = QLabel("所有歌曲已处理完成！")
            no_issues_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            no_issues_label.setStyleSheet("font-size: 12pt; color: green; padding: 20px;")
            self.scroll_layout.addWidget(no_issues_label)
        
        # 重新添加stretch
        self.scroll_layout.addStretch()
    
    def set_all_checked(self, checked):
        for i in range(self.scroll_layout.count()):
            widget = self.scroll_layout.itemAt(i).widget()
            if widget and hasattr(widget, 'checkbox'):
                widget.checkbox.setChecked(checked)

    def toggle_filter(self):
        if not self.filter_shown:
            for i in range(self.scroll_layout.count()):
                widget = self.scroll_layout.itemAt(i).widget()
                if widget and hasattr(widget, 'checkbox'):
                    song = self.songs_data[i]
                    if song["status"] not in ("ok", "skipped"):
                        widget.setVisible(True)
                    else:
                        widget.setVisible(False)
            self.filter_btn.setText("显示全部")
            self.filter_shown = True
        else:
            for i in range(self.scroll_layout.count()):
                widget = self.scroll_layout.itemAt(i).widget()
                if widget:
                    widget.setVisible(True)
            self.filter_btn.setText("仅显示需确认")
            self.filter_shown = False

    def confirm_only(self):
        logger.info("用户确认，开始生成文档")
        
        selected_songs = []
        unselected_songs = []
        
        for i, song in enumerate(self.songs_data):
            if song["status"] == "ok":
                selected_songs.append(song)
            elif song["status"] == "skipped":
                unselected_songs.append(song)
            elif song["status"] == "pending":
                checkbox = self.find_checkbox_for_song(i)
                if checkbox and checkbox.isChecked():
                    song["status"] = "ok"
                    song["check_message"] = "用户确认包含"
                    selected_songs.append(song)
                else:
                    song["status"] = "skipped"
                    song["check_message"] = "用户跳过"
                    unselected_songs.append(song)
            elif song["status"] == "warning" or song["status"] == "error":
                checkbox = self.find_checkbox_for_song(i)
                if checkbox and checkbox.isChecked():
                    selected_songs.append(song)
                else:
                    unselected_songs.append(song)
        
        if not selected_songs:
            QMessageBox.warning(self, "警告", "请至少选择一首歌曲！")
            return
        
        try:
            self.generate_word(selected_songs, unselected_songs)
            QMessageBox.information(self, "完成", f"歌词已保存至:\n{self.output_file}")
            self.accept()
        except Exception as e:
            QMessageBox.critical(self, "错误", f"生成 Word 文档失败:\n{e}")

    def confirm_and_generate(self):
        # 检查是否有正在搜索的歌曲
        searching_songs = [s for s in self.songs_data if s["status"] == "searching"]
        if searching_songs:
            QMessageBox.warning(self, "警告", "有歌曲正在搜索中，请等待完成后再生成文档")
            return
        
        selected_songs = []
        unselected_songs = []
        
        for i, song in enumerate(self.songs_data):
            if song["status"] == "ok":
                selected_songs.append(song)
            elif song["status"] == "skipped":
                unselected_songs.append(song)
            elif song["status"] == "pending":
                checkbox = self.find_checkbox_for_song(i)
                if checkbox and checkbox.isChecked():
                    song["status"] = "ok"
                    song["check_message"] = "用户确认包含"
                    selected_songs.append(song)
                else:
                    song["status"] = "skipped"
                    song["check_message"] = "用户跳过"
                    unselected_songs.append(song)
            elif song["status"] == "warning" or song["status"] == "error":
                checkbox = self.find_checkbox_for_song(i)
                if checkbox and checkbox.isChecked():
                    song["status"] = "ok"
                    song["check_message"] = "用户确认包含"
                    selected_songs.append(song)
                else:
                    song["status"] = "skipped"
                    song["check_message"] = "用户跳过"
                    unselected_songs.append(song)
        
        if not selected_songs:
            QMessageBox.warning(self, "警告", "请至少选择一首歌曲！")
            return
        
        try:
            self.generate_word(selected_songs, unselected_songs)
            QMessageBox.information(self, "完成", f"歌词已保存至:\n{self.output_file}")
            self.accept()
        except Exception as e:
            QMessageBox.critical(self, "错误", f"生成 Word 文档失败:\n{e}")

    def find_checkbox_for_song(self, song_index):
        for i in range(self.scroll_layout.count()):
            widget = self.scroll_layout.itemAt(i).widget()
            if widget and hasattr(widget, 'checkbox') and hasattr(widget.checkbox, 'song_index'):
                if widget.checkbox.song_index == song_index:
                    return widget.checkbox
        return None

    def generate_word(self, songs_data, unselected_songs=None):
        if unselected_songs is None:
            unselected_songs = []
        
        logger.info(f"开始生成Word文档, 选中歌曲: {len(songs_data)}, 未选中: {len(unselected_songs)}")
        
        doc = Document()
        
        title_paragraph = doc.add_heading('歌词合集', 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 收集未找到歌词的歌曲和跳过的歌曲
        missing_songs = [s for s in songs_data if not s.get('lyrics')]
        skipped_songs = unselected_songs
        successful_songs = [s for s in songs_data if s.get('lyrics') and s not in unselected_songs]
        
        logger.info(f"成功获取歌词: {len(successful_songs)}, 未找到: {len(missing_songs)}, 跳过: {len(skipped_songs)}")
        
        # 未找到歌词的歌曲列表 - 放第一页
        if missing_songs:
            logger.warning(f"未找到歌词的歌曲: {[s['matched_title'] for s in missing_songs]}")
            doc.add_paragraph()
            warning_heading = doc.add_heading('未找到歌词的歌曲', level=2)
            for song in missing_songs:
                doc.add_paragraph(f"• {song['matched_title']} - {song['matched_artist']}")
        
        # 跳过的歌曲列表 - 放第一页
        if skipped_songs:
            logger.info(f"用户跳过的歌曲: {[s['input_title'] for s in skipped_songs]}")
            doc.add_paragraph()
            skipped_heading = doc.add_heading('跳过的歌曲', level=2)
            for song in skipped_songs:
                doc.add_paragraph(f"• {song['input_title']} - {song['input_artist']}")
        
        doc.add_paragraph()
        doc.add_paragraph("目录")
        doc.add_paragraph()
        
        for i, song in enumerate(successful_songs, 1):
            doc.add_paragraph(f"{i}. {song['matched_title']} - {song['matched_artist']}")
        
        for song in successful_songs:
            doc.add_page_break()
            
            heading = doc.add_heading(f"{song['matched_title']} - {song['matched_artist']}", level=1)
            
            if song['lyrics']:
                para = doc.add_paragraph(song['lyrics'])
                for run in para.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
            else:
                doc.add_paragraph("（未找到歌词）")
        
        os.makedirs(os.path.dirname(self.output_file), exist_ok=True)
        doc.save(self.output_file)
        
        logger.info(f"Word文档生成成功: {self.output_file}")


class LyricsGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        logger.info("GUI 初始化")
        self.download_thread = None
        self.songs_data = []
        self.init_ui()
        logger.info("GUI 初始化完成")

    def init_ui(self):
        self.setWindowTitle("歌词批量下载工具 - Word 文档生成器")
        self.setMinimumSize(700, 600)
        
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QVBoxLayout()
        central_widget.setLayout(main_layout)
        
        title_label = QLabel("歌词批量下载工具 - Word 文档生成器")
        title_font = QFont()
        title_font.setPointSize(14)
        title_font.setBold(True)
        title_label.setFont(title_font)
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        main_layout.addWidget(title_label)
        
        main_layout.addSpacing(10)
        
        input_group = QGroupBox("文件设置")
        input_layout = QVBoxLayout()
        
        file_layout1 = QHBoxLayout()
        file_layout1.addWidget(QLabel("歌名列表文件:"))
        self.input_file_edit = QLineEdit()
        self.input_file_edit.setPlaceholderText("选择 songs.txt 文件...")
        file_layout1.addWidget(self.input_file_edit)
        self.input_file_btn = QPushButton("浏览...")
        self.input_file_btn.clicked.connect(self.select_input_file)
        file_layout1.addWidget(self.input_file_btn)
        input_layout.addLayout(file_layout1)
        
        file_layout2 = QHBoxLayout()
        file_layout2.addWidget(QLabel("输出文件路径:"))
        self.output_file_edit = QLineEdit()
        self.output_file_edit.setPlaceholderText("选择输出 .docx 文件路径...")
        self.output_file_edit.setText("output/歌词合集.docx")
        file_layout2.addWidget(self.output_file_edit)
        self.output_file_btn = QPushButton("浏览...")
        self.output_file_btn.clicked.connect(self.select_output_file)
        file_layout2.addWidget(self.output_file_btn)
        input_layout.addLayout(file_layout2)
        
        input_group.setLayout(input_layout)
        main_layout.addWidget(input_group)
        
        progress_group = QGroupBox("下载进度")
        progress_layout = QVBoxLayout()
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        progress_layout.addWidget(self.progress_bar)
        
        self.current_song_label = QLabel("等待开始...")
        self.current_song_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        progress_layout.addWidget(self.current_song_label)
        
        progress_group.setLayout(progress_layout)
        main_layout.addWidget(progress_group)
        
        log_group = QGroupBox("日志")
        log_layout = QVBoxLayout()
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMaximumHeight(200)
        log_layout.addWidget(self.log_text)
        log_group.setLayout(log_layout)
        main_layout.addWidget(log_group)
        
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        self.start_btn = QPushButton("开始搜索")
        self.start_btn.setMinimumHeight(40)
        self.start_btn.setStyleSheet("QPushButton { background-color: #4CAF50; color: white; font-size: 12pt; }")
        self.start_btn.clicked.connect(self.start_download)
        button_layout.addWidget(self.start_btn)
        
        self.stop_btn = QPushButton("停止")
        self.stop_btn.setMinimumHeight(40)
        self.stop_btn.setEnabled(False)
        self.stop_btn.setStyleSheet("QPushButton { background-color: #f44336; color: white; font-size: 12pt; }")
        self.stop_btn.clicked.connect(self.stop_download)
        button_layout.addWidget(self.stop_btn)
        button_layout.addStretch()
        
        main_layout.addLayout(button_layout)

    def select_input_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择歌名列表文件", "", "文本文件 (*.txt);;所有文件 (*.*)"
        )
        if file_path:
            self.input_file_edit.setText(file_path)

    def select_output_file(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "选择输出文件", "歌词合集.docx", "Word 文档 (*.docx)"
        )
        if file_path:
            self.output_file_edit.setText(file_path)

    def start_download(self):
        logger.info("用户点击开始下载按钮")
        
        input_file = self.input_file_edit.text().strip()
        output_file = self.output_file_edit.text().strip()
        
        logger.info(f"输入文件: {input_file}")
        logger.info(f"输出文件: {output_file}")
        
        if not input_file:
            logger.warning("未选择输入文件")
            QMessageBox.warning(self, "警告", "请选择歌名列表文件！")
            return
        
        if not output_file:
            logger.warning("未选择输出文件")
            QMessageBox.warning(self, "警告", "请选择输出文件路径！")
            return
        
        self.start_btn.setEnabled(False)
        self.stop_btn.setEnabled(True)
        self.progress_bar.setValue(0)
        self.log_text.clear()
        
        logger.info("启动下载线程...")
        self.download_thread = LyricsDownloaderThread(input_file, output_file)
        self.download_thread.progress_updated.connect(self.update_progress)
        self.download_thread.log_message.connect(self.append_log)
        self.download_thread.finished_with_result.connect(self.download_finished)
        self.download_thread.stopped.connect(self.download_stopped)
        self.download_thread.start()
        
        logger.info("下载线程已启动")
    
    def stop_download(self):
        if self.download_thread:
            self.download_thread.stop()
    
    def update_progress(self, current, total, song_name):
        percentage = int(current / total * 100)
        self.progress_bar.setValue(percentage)
        self.current_song_label.setText(f"正在获取: {song_name} ({current}/{total})")

    def append_log(self, message):
        self.log_text.append(message)

    def download_finished(self, songs_data):
        self.songs_data = songs_data
        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        self.current_song_label.setText("下载完成！等待确认...")
        
        output_file = self.output_file_edit.text().strip()
        
        dialog = PreviewDialog(songs_data, output_file, self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.log_text.append("\n=== Word 文档已生成 ===")
        else:
            self.log_text.append("\n=== 用户取消 ===")

    def download_stopped(self):
        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        self.current_song_label.setText("已停止")


def main():
    app = QApplication(sys.argv)
    window = LyricsGUI()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()